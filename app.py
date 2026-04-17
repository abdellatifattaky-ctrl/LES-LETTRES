import streamlit as st
import google.generativeai as genai
import io
import sqlite3
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- إعدادات الصفحة ---
st.set_page_config(page_title="نظام المراسلات - جماعة أسكاون", layout="wide")

# --- إعداد الذكاء الاصطناعي ---
MY_API_KEY = "AQ.Ab8RN6LtDoih_ytIju3ulJTIa18hvJdnOpfnAUpn3KMJVDNb1w"
genai.configure(api_key=MY_API_KEY)

@st.cache_resource
def get_available_model():
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except: return "models/gemini-1.5-flash"

active_model_name = get_available_model()

# --- وظائف قاعدة البيانات ---
def init_db():
    conn = sqlite3.connect('askaouen_admin.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS letters 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  letter_number TEXT, date TEXT, sender TEXT, 
                  recipient TEXT, subject TEXT, content TEXT, attachments TEXT)''')
    conn.commit()
    conn.close()

def get_next_num():
    try:
        conn = sqlite3.connect('askaouen_admin.db')
        c = conn.cursor()
        c.execute("SELECT COUNT(*) FROM letters")
        count = c.fetchone()[0]
        conn.close()
        return f"{count + 7:02d}/{datetime.now().year}" # يبدأ من الرقم 07 حسب ملفك
    except: return f"07/{datetime.now().year}"

# --- دالة تنسيق ملف Word (نموذج جماعة أسكاون) ---
def create_askaouen_word(l_num, date_str, sender, recipient, subject, content, attachments):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # 1. الرأس الإداري (يمين ويسار) 
    header_table = doc.add_table(rows=1, cols=2)
    header_table.width = Inches(6)
    
    right_cell = header_table.cell(0, 0)
    p_right = right_cell.add_paragraph()
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # معلومات مأخوذة من ترويسة ملفك 
    p_right.add_run("المملكة المغربية\nوزارة الداخلية\nإقليم تارودانت\nدائرة تالوين\nجماعة أسكاون").bold = True
    
    left_cell = header_table.cell(0, 1)
    p_left = left_cell.add_paragraph()
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_left.add_run(f"أسكاون في: {date_str}") [cite: 14]

    doc.add_paragraph()

    # 2. الرقم الترتيبي [cite: 13]
    ref_p = doc.add_paragraph()
    ref_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ref_p.add_run(f"عدد: {l_num}").bold = True

    doc.add_paragraph()

    # 3. من ... إلى [cite: 15, 16, 17]
    dest_p = doc.add_paragraph()
    dest_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dest_p.add_run(f"{sender}\nإلى\n{recipient}").bold = True

    doc.add_paragraph()

    # 4. الموضوع [cite: 18]
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sub_p.add_run(f"الموضوع: {subject}").bold = True
    
    # 5. المرفقات (إضافة جديدة)
    if attachments:
        att_p = doc.add_paragraph()
        att_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        att_p.add_run(f"المرفقات: {attachments}")

    doc.add_paragraph()

    # 6. التحية الرسمية [cite: 19]
    salutation = doc.add_paragraph()
    salutation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    salutation.add_run("سلام تام بوجود مولانا الإمام").bold = True
    
    # 7. نص المراسلة 
    body_p = doc.add_paragraph()
    body_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    body_p.add_run("وبعد، " + content)
    body_p.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()

    # 8. الخاتمة والتوقيع [cite: 21, 22]
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    end_p.add_run("وتقبلوا أسمى عبارات التقدير والاحترام").bold = True
    
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sign_p.add_run("\nرئيس المجلس الجماعي").bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# --- واجهة المستخدم Streamlit ---
init_db()
st.title("🏛️ بوابة المراسلات الإدارية - جماعة أسكاون")

col1, col2 = st.columns([1, 1.5])

with col1:
    st.subheader("🖋️ تحرير مراسلة")
    l_num = st.text_input("رقم المراسلة", value=get_next_num())
    date_val = st.text_input("التاريخ", value=datetime.now().strftime("%d %B %Y"))
    sender = st.text_input("من", value="رئيس جماعة أسكاون") [cite: 15]
    recipient = st.text_input("إلى", value="السيد القابض الجماعي بتالوين") [cite: 17]
    subject = st.text_input("الموضوع", value="حضور جلسة فتح الاظرفة") [cite: 18]
    attachments = st.text_input("المرفقات (إن وجدت)")
    hint = st.text_area("عناصر المحتوى (مثال: أشغال تزويد الماء بدواوير...)") [cite: 20]
    
    generate_btn = st.button("🚀 صياغة المراسلة")

with col2:
    st.subheader("📄 المعاينة")
    if generate_btn:
        with st.spinner("جاري التوليد..."):
            model = genai.GenerativeModel(active_model_name)
            prompt = f"صغ خطابا إداريا مغربيا لجماعة أسكاون. الموضوع: {subject}. العناصر: {hint}. استخدم عبارات: يشرفني دعوتكم، علاقة بالموضوع، يوم الاربعاء 21 يناير."
            response = model.generate_content(prompt)
            st.session_state['aska_text'] = response.text

    if 'aska_text' in st.session_state:
        final_text = st.text_area("النص المولد:", value=st.session_state['aska_text'], height=400)
        
        if st.button("💾 حفظ وتنزيل"):
            conn = sqlite3.connect('askaouen_admin.db')
            c = conn.cursor()
            c.execute("INSERT INTO letters (letter_number, date, sender, recipient, subject, content, attachments) VALUES (?,?,?,?,?,?,?)",
                      (l_num, date_val, sender, recipient, subject, final_text, attachments))
            conn.commit()
            conn.close()
            
            word_file = create_askaouen_word(l_num, date_val, sender, recipient, subject, final_text, attachments)
            st.download_button("تحميل المراسلة الرسمية (.docx)", word_file, f"askaouen_{l_num.replace('/', '-')}.docx")

st.markdown("---")
st.subheader("🗄️ أرشيف الجماعة")
conn = sqlite3.connect('askaouen_admin.db')
df = pd.read_sql_query("SELECT * FROM letters ORDER BY id DESC", conn)
conn.close()
st.dataframe(df, use_container_width=True)
